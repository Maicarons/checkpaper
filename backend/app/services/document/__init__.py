"""
CheckPaper 文档服务模块
"""
import os
import uuid
import json
from typing import List, Optional
from datetime import datetime
from sqlmodel import Session, select, func

from ...schemas.document import (
    Document,
    DocumentTypeEnum,
    DocumentStatusEnum,
    ParsedContent
)
from ...core.config import settings


class DocumentService:
    """文档服务类"""
    
    def create_document(self, session: Session, document_data: dict) -> Document:
        """
        创建文档记录
        
        Args:
            session: 数据库会话
            document_data: 文档数据
        
        Returns:
            创建的文档对象
        """
        document = Document(**document_data)
        session.add(document)
        session.commit()
        session.refresh(document)
        return document
    
    def get_document(self, session: Session, document_id: str) -> Optional[Document]:
        """
        获取文档
        
        Args:
            session: 数据库会话
            document_id: 文档ID
        
        Returns:
            文档对象或 None
        """
        statement = select(Document).where(Document.id == document_id)
        return session.exec(statement).first()
    
    def get_documents(
        self,
        session: Session,
        offset: int = 0,
        limit: int = 20,
        status: Optional[str] = None,
        document_type: Optional[DocumentTypeEnum] = None
    ) -> List[Document]:
        """
        获取文档列表
        
        Args:
            session: 数据库会话
            offset: 偏移量
            limit: 限制数量
            status: 状态过滤
            document_type: 文档类型过滤
        
        Returns:
            文档列表
        """
        statement = select(Document)
        
        if status:
            statement = statement.where(Document.status == status)
        if document_type:
            statement = statement.where(Document.file_type == document_type)
        
        statement = statement.offset(offset).limit(limit).order_by(Document.upload_time.desc())
        return list(session.exec(statement).all())
    
    def get_documents_count(
        self,
        session: Session,
        status: Optional[str] = None,
        document_type: Optional[DocumentTypeEnum] = None
    ) -> int:
        """
        获取文档总数
        
        Args:
            session: 数据库会话
            status: 状态过滤
            document_type: 文档类型过滤
        
        Returns:
            文档总数
        """
        statement = select(func.count()).select_from(Document)
        
        if status:
            statement = statement.where(Document.status == status)
        if document_type:
            statement = statement.where(Document.file_type == document_type)
        
        return session.exec(statement).one()
    
    def update_document_status(
        self,
        session: Session,
        document_id: str,
        status: DocumentStatusEnum,
        error_message: Optional[str] = None
    ) -> Optional[Document]:
        """
        更新文档状态
        
        Args:
            session: 数据库会话
            document_id: 文档ID
            status: 新状态
            error_message: 错误信息
        
        Returns:
            更新后的文档对象
        """
        document = self.get_document(session, document_id)
        if not document:
            return None
        
        document.status = status
        if error_message:
            document.error_message = error_message
        if status == DocumentStatusEnum.PARSED:
            document.parsed_time = datetime.utcnow()
        
        session.add(document)
        session.commit()
        session.refresh(document)
        return document
    
    def delete_document(self, session: Session, document_id: str) -> bool:
        """
        删除文档
        
        Args:
            session: 数据库会话
            document_id: 文档ID
        
        Returns:
            是否删除成功
        """
        document = self.get_document(session, document_id)
        if not document:
            return False
        
        session.delete(document)
        session.commit()
        return True
    
    def parse_document(self, session: Session, document_id: str) -> ParsedContent:
        """
        解析文档
        
        Args:
            session: 数据库会话
            document_id: 文档ID
        
        Returns:
            解析后的内容
        """
        document = self.get_document(session, document_id)
        if not document:
            raise ValueError("文档不存在")
        
        # 更新状态为解析中
        self.update_document_status(session, document_id, DocumentStatusEnum.PARSING)
        
        try:
            # 根据文档类型选择解析器
            if document.file_type == DocumentTypeEnum.PDF:
                content = self._parse_pdf(document.file_path)
            elif document.file_type == DocumentTypeEnum.WORD:
                content = self._parse_word(document.file_path)
            elif document.file_type == DocumentTypeEnum.LATEX:
                content = self._parse_latex(document.file_path)
            else:
                raise ValueError(f"不支持的文档类型: {document.file_type}")
            
            # 更新状态为已解析
            self.update_document_status(session, document_id, DocumentStatusEnum.PARSED)
            
            return content
            
        except Exception as e:
            # 更新状态为失败
            self.update_document_status(
                session, document_id, DocumentStatusEnum.FAILED, str(e)
            )
            raise
    
    def _parse_pdf(self, file_path: str) -> ParsedContent:
        """
        解析 PDF 文档
        
        Args:
            file_path: 文件路径
        
        Returns:
            解析后的内容
        """
        import pymupdf
        
        doc = pymupdf.open(file_path)
        text_parts = []
        figures = []
        tables = []
        
        for page_num, page in enumerate(doc):
            # 提取文本
            text_parts.append(page.get_text())
            
            # 提取图片信息
            image_list = page.get_images()
            for img_index, img in enumerate(image_list):
                figures.append({
                    "page": page_num + 1,
                    "index": img_index,
                    "xref": img[0]
                })
        
        full_text = "\n".join(text_parts)
        
        # 简单的标题提取（第一页的第一行）
        title = text_parts[0].split("\n")[0].strip() if text_parts else ""
        
        doc.close()
        
        return ParsedContent(
            title=title,
            raw_text=full_text,
            figures=figures,
            tables=tables,
            sections=[],
            references=[],
            citations=[],
            metadata={"page_count": len(text_parts)}
        )
    
    def _parse_word(self, file_path: str) -> ParsedContent:
        """
        解析 Word 文档
        
        Args:
            file_path: 文件路径
        
        Returns:
            解析后的内容
        """
        from docx import Document as DocxDocument
        
        doc = DocxDocument(file_path)
        
        paragraphs = []
        tables = []
        
        for para in doc.paragraphs:
            paragraphs.append(para.text)
        
        for table in doc.tables:
            table_data = []
            for row in table.rows:
                row_data = [cell.text for cell in row.cells]
                table_data.append(row_data)
            tables.append(table_data)
        
        full_text = "\n".join(paragraphs)
        title = paragraphs[0] if paragraphs else ""
        
        return ParsedContent(
            title=title,
            raw_text=full_text,
            figures=[],
            tables=tables,
            sections=[],
            references=[],
            citations=[],
            metadata={"paragraph_count": len(paragraphs)}
        )
    
    def _parse_latex(self, file_path: str) -> ParsedContent:
        """
        解析 LaTeX 文档
        
        Args:
            file_path: 文件路径
        
        Returns:
            解析后的内容
        """
        with open(file_path, 'r', encoding='utf-8') as f:
            content = f.read()
        
        # 简单的 LaTeX 解析
        # 提取标题
        import re
        title_match = re.search(r'\\title\{([^}]+)\}', content)
        title = title_match.group(1) if title_match else ""
        
        # 提取摘要
        abstract_match = re.search(r'\\begin\{abstract\}(.+?)\\end\{abstract\}', content, re.DOTALL)
        abstract = abstract_match.group(1).strip() if abstract_match else ""
        
        # 提取引用
        citations = re.findall(r'\\cite\{([^}]+)\}', content)
        citation_list = []
        for cite in citations:
            for c in cite.split(','):
                citation_list.append({"key": c.strip()})
        
        return ParsedContent(
            title=title,
            abstract=abstract,
            raw_text=content,
            figures=[],
            tables=[],
            sections=[],
            references=[],
            citations=citation_list,
            metadata={"file_size": os.path.getsize(file_path)}
        )
