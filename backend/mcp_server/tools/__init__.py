"""
CheckPaper MCP 工具实现
"""
import os
import json
import re
from typing import Dict, Any, List, Optional
import httpx


async def parse_pdf(file_path: str) -> Dict[str, Any]:
    """
    解析PDF文件
    
    Args:
        file_path: PDF文件路径
    
    Returns:
        解析结果
    """
    try:
        import pymupdf
        
        if not os.path.exists(file_path):
            return {"error": f"文件不存在: {file_path}"}
        
        doc = pymupdf.open(file_path)
        text_parts = []
        figures = []
        
        for page_num, page in enumerate(doc):
            text_parts.append(page.get_text())
            
            # 提取图片信息
            image_list = page.get_images()
            for img_index, img in enumerate(image_list):
                figures.append({
                    "page": page_num + 1,
                    "index": img_index,
                    "xref": img[0]
                })
        
        full_text = "\n".join(text_parts)
        
        # 提取标题（假设第一行是标题）
        title = text_parts[0].split("\n")[0].strip() if text_parts else ""
        
        # 提取参考文献
        references = _extract_references(full_text)
        
        # 提取引用
        citations = _extract_citations(full_text)
        
        doc.close()
        
        return {
            "title": title,
            "content": full_text,
            "figures": figures,
            "references": references,
            "citations": citations,
            "page_count": len(text_parts)
        }
    
    except ImportError:
        return {"error": "需要安装 pymupdf 库"}
    except Exception as e:
        return {"error": str(e)}


async def parse_docx(file_path: str) -> Dict[str, Any]:
    """
    解析Word文档
    
    Args:
        file_path: Word文件路径
    
    Returns:
        解析结果
    """
    try:
        from docx import Document
        
        if not os.path.exists(file_path):
            return {"error": f"文件不存在: {file_path}"}
        
        doc = Document(file_path)
        
        paragraphs = []
        tables = []
        
        for para in doc.paragraphs:
            paragraphs.append(para.text)
        
        for table in doc.tables:
            table_data = []
            for row in table.rows:
                row_data = [cell.text for cell in row.cells]
                table_data.append(row_data)
            tables.append(table_data)
        
        full_text = "\n".join(paragraphs)
        title = paragraphs[0] if paragraphs else ""
        
        return {
            "title": title,
            "content": full_text,
            "tables": tables,
            "paragraph_count": len(paragraphs)
        }
    
    except ImportError:
        return {"error": "需要安装 python-docx 库"}
    except Exception as e:
        return {"error": str(e)}


async def parse_latex(file_path: str) -> Dict[str, Any]:
    """
    解析LaTeX文档
    
    Args:
        file_path: LaTeX文件路径
    
    Returns:
        解析结果
    """
    try:
        if not os.path.exists(file_path):
            return {"error": f"文件不存在: {file_path}"}
        
        with open(file_path, 'r', encoding='utf-8') as f:
            content = f.read()
        
        # 提取标题
        title_match = re.search(r'\\title\{([^}]+)\}', content)
        title = title_match.group(1) if title_match else ""
        
        # 提取摘要
        abstract_match = re.search(r'\\begin\{abstract\}(.+?)\\end\{abstract\}', content, re.DOTALL)
        abstract = abstract_match.group(1).strip() if abstract_match else ""
        
        # 提取引用
        citations = re.findall(r'\\cite\{([^}]+)\}', content)
        citation_list = []
        for cite in citations:
            for c in cite.split(','):
                citation_list.append({"key": c.strip()})
        
        # 提取参考文献（从 .bib 文件）
        references = []
        bib_match = re.search(r'\\bibliography\{([^}]+)\}', content)
        if bib_match:
            bib_file = bib_match.group(1) + ".bib"
            bib_path = os.path.join(os.path.dirname(file_path), bib_file)
            if os.path.exists(bib_path):
                references = _parse_bib_file(bib_path)
        
        return {
            "title": title,
            "abstract": abstract,
            "content": content,
            "citations": citation_list,
            "references": references,
            "file_size": len(content)
        }
    
    except Exception as e:
        return {"error": str(e)}


async def check_citations(
    content: str,
    references: List[Dict],
    citations: List[Dict]
) -> Dict[str, Any]:
    """
    检查引用一致性
    
    Args:
        content: 论文正文
        references: 参考文献列表
        citations: 引用列表
    
    Returns:
        检查结果
    """
    issues = []
    
    # 提取参考文献的键
    ref_keys = set()
    for ref in references:
        if "key" in ref:
            ref_keys.add(ref["key"])
        elif "id" in ref:
            ref_keys.add(ref["id"])
    
    # 检查每个引用
    cited_keys = set()
    for cite in citations:
        key = cite.get("key", cite.get("id", ""))
        cited_keys.add(key)
        
        if key not in ref_keys:
            issues.append({
                "severity": "critical",
                "category": "引用错误",
                "title": f"引用 '{key}' 在参考文献中不存在",
                "description": f"正文中引用了 '{key}'，但在参考文献列表中找不到对应条目",
                "suggestion": "检查引用键是否正确，或添加缺失的参考文献"
            })
    
    # 检查未引用的参考文献
    for ref in references:
        key = ref.get("key", ref.get("id", ""))
        if key not in cited_keys:
            issues.append({
                "severity": "warning",
                "category": "未引用文献",
                "title": f"参考文献 '{key}' 未被引用",
                "description": f"参考文献列表中的 '{key}' 在正文中未被引用",
                "suggestion": "删除未使用的参考文献或在正文中添加引用"
            })
    
    return {
        "issues": issues,
        "total_references": len(references),
        "total_citations": len(citations),
        "matched": len(ref_keys.intersection(cited_keys))
    }


async def verify_references(references: List[Dict]) -> Dict[str, Any]:
    """
    验证参考文献真实性
    
    Args:
        references: 参考文献列表
    
    Returns:
        验证结果
    """
    results = []
    
    for ref in references:
        doi = ref.get("doi", "")
        title = ref.get("title", "")
        
        verified = False
        confidence = 0.0
        issues = []
        
        # 通过 DOI 验证
        if doi:
            try:
                async with httpx.AsyncClient() as client:
                    response = await client.get(
                        f"https://api.crossref.org/works/{doi}",
                        timeout=10.0
                    )
                    if response.status_code == 200:
                        data = response.json()
                        verified = True
                        confidence = 0.9
                        
                        # 验证标题匹配
                        if title:
                            crossref_title = data.get("message", {}).get("title", [""])[0]
                            if title.lower() not in crossref_title.lower():
                                issues.append("标题不匹配")
                                confidence -= 0.2
            except Exception as e:
                issues.append(f"DOI验证失败: {str(e)}")
        
        # 通过标题搜索验证
        elif title:
            try:
                async with httpx.AsyncClient() as client:
                    response = await client.get(
                        "https://api.crossref.org/works",
                        params={"query.title": title, "rows": 1},
                        timeout=10.0
                    )
                    if response.status_code == 200:
                        data = response.json()
                        items = data.get("message", {}).get("items", [])
                        if items:
                            verified = True
                            confidence = 0.7
            except Exception as e:
                issues.append(f"标题搜索失败: {str(e)}")
        
        results.append({
            "reference": ref,
            "verified": verified,
            "confidence": confidence,
            "issues": issues
        })
    
    return {
        "results": results,
        "total": len(references),
        "verified": sum(1 for r in results if r["verified"])
    }


async def check_figures(
    content: str,
    figures: List[Dict],
    tables: List[Dict]
) -> Dict[str, Any]:
    """
    检查图表引用完整性
    
    Args:
        content: 论文正文
        figures: 图片列表
        tables: 表格列表
    
    Returns:
        检查结果
    """
    issues = []
    
    # 提取正文中对图表的引用
    figure_refs = set(re.findall(r'(?:Fig(?:ure)?\.?\s*(\d+)|图\s*(\d+))', content, re.IGNORECASE))
    table_refs = set(re.findall(r'(?:Table\.?\s*(\d+)|表\s*(\d+))', content, re.IGNORECASE))
    
    # 扁平化引用集合
    figure_refs = {ref[0] or ref[1] for ref in figure_refs}
    table_refs = {ref[0] or ref[1] for ref in table_refs}
    
    # 检查图片引用
    for fig in figures:
        fig_id = str(fig.get("id", fig.get("index", "")))
        if fig_id and fig_id not in figure_refs:
            issues.append({
                "severity": "warning",
                "category": "未引用图片",
                "title": f"图片 {fig_id} 未被引用",
                "description": f"第{fig.get('page', '?')}页的图片在正文中未被引用",
                "suggestion": "在正文中添加对图片的引用或删除未使用的图片"
            })
    
    # 检查表格引用
    for tab in tables:
        tab_id = str(tab.get("id", ""))
        if tab_id and tab_id not in table_refs:
            issues.append({
                "severity": "warning",
                "category": "未引用表格",
                "title": f"表格 {tab_id} 未被引用",
                "description": "表格在正文中未被引用",
                "suggestion": "在正文中添加对表格的引用或删除未使用的表格"
            })
    
    return {
        "issues": issues,
        "figures_count": len(figures),
        "tables_count": len(tables),
        "figure_refs_count": len(figure_refs),
        "table_refs_count": len(table_refs)
    }


async def check_format(content: str, metadata: Dict = None) -> Dict[str, Any]:
    """
    检查论文格式
    
    Args:
        content: 论文内容
        metadata: 文档元数据
    
    Returns:
        检查结果
    """
    issues = []
    
    # 检查基本结构
    sections = ["摘要", "引言", "方法", "结果", "讨论", "结论"]
    found_sections = []
    
    for section in sections:
        if section in content:
            found_sections.append(section)
    
    if len(found_sections) < 3:
        issues.append({
            "severity": "warning",
            "category": "结构不完整",
            "title": "论文章节结构不完整",
            "description": f"只找到以下章节: {', '.join(found_sections)}",
            "suggestion": "确保论文包含完整的章节结构"
        })
    
    # 检查图表编号
    fig_numbers = re.findall(r'(?:Fig(?:ure)?\.?\s*(\d+)|图\s*(\d+))', content, re.IGNORECASE)
    if fig_numbers:
        nums = sorted(set(int(n[0] or n[1]) for n in fig_numbers))
        expected = list(range(1, len(nums) + 1))
        if nums != expected:
            issues.append({
                "severity": "info",
                "category": "图表编号",
                "title": "图表编号可能不连续",
                "description": f"发现的图表编号: {nums}",
                "suggestion": "检查图表编号是否连续"
            })
    
    return {
        "issues": issues,
        "sections_found": found_sections,
        "content_length": len(content)
    }


async def web_search_reference(
    title: str,
    authors: str = "",
    doi: str = ""
) -> Dict[str, Any]:
    """
    联网搜索验证参考文献
    
    Args:
        title: 文献标题
        authors: 作者
        doi: DOI号
    
    Returns:
        搜索结果
    """
    try:
        async with httpx.AsyncClient() as client:
            # 通过 Crossref 搜索
            params = {"query.title": title, "rows": 3}
            if authors:
                params["query.author"] = authors
            
            response = await client.get(
                "https://api.crossref.org/works",
                params=params,
                timeout=15.0
            )
            
            if response.status_code == 200:
                data = response.json()
                items = data.get("message", {}).get("items", [])
                
                results = []
                for item in items[:3]:
                    results.append({
                        "title": item.get("title", [""])[0],
                        "doi": item.get("DOI", ""),
                        "authors": [
                            f"{a.get('given', '')} {a.get('family', '')}"
                            for a in item.get("author", [])
                        ],
                        "journal": item.get("container-title", [""])[0],
                        "year": item.get("published", {}).get("date-parts", [[None]])[0][0],
                        "url": item.get("URL", "")
                    })
                
                return {
                    "found": len(results) > 0,
                    "results": results,
                    "query": title
                }
            
            return {"found": False, "error": "搜索失败"}
    
    except Exception as e:
        return {"found": False, "error": str(e)}


def _extract_references(text: str) -> List[Dict]:
    """提取参考文献"""
    references = []
    
    # 查找参考文献部分
    ref_pattern = r'(?:References|参考文献|Bibliography)\s*\n([\s\S]+?)(?:\n\n|\Z)'
    ref_match = re.search(ref_pattern, text, re.IGNORECASE)
    
    if ref_match:
        ref_text = ref_match.group(1)
        # 分割每条参考文献
        ref_lines = re.split(r'\n(?=\[\d+\]|\d+\.)', ref_text)
        
        for i, line in enumerate(ref_lines):
            line = line.strip()
            if line:
                # 提取编号
                num_match = re.match(r'\[(\d+)\]', line)
                ref_id = num_match.group(1) if num_match else str(i + 1)
                
                # 提取 DOI
                doi_match = re.search(r'doi:\s*(10\.\S+)', line, re.IGNORECASE)
                doi = doi_match.group(1) if doi_match else ""
                
                references.append({
                    "id": ref_id,
                    "text": line,
                    "doi": doi
                })
    
    return references


def _extract_citations(text: str) -> List[Dict]:
    """提取正文引用"""
    citations = []
    
    # 匹配 [1], [2,3], [1-5] 格式
    bracket_refs = re.findall(r'\[(\d+(?:[,\s]*\d+)*(?:\s*-\s*\d+)?)\]', text)
    
    for ref in bracket_refs:
        # 处理范围引用如 1-5
        if '-' in ref:
            parts = ref.split('-')
            try:
                start, end = int(parts[0].strip()), int(parts[1].strip())
                for i in range(start, end + 1):
                    citations.append({"key": str(i)})
            except ValueError:
                citations.append({"key": ref})
        else:
            # 处理逗号分隔的引用如 1,2,3
            for num in re.findall(r'\d+', ref):
                citations.append({"key": num})
    
    # 匹配 \cite{key1, key2} 格式
    cite_refs = re.findall(r'\\cite\{([^}]+)\}', text)
    for ref in cite_refs:
        for key in ref.split(','):
            citations.append({"key": key.strip()})
    
    return citations


def _parse_bib_file(file_path: str) -> List[Dict]:
    """解析 .bib 文件"""
    references = []
    
    try:
        with open(file_path, 'r', encoding='utf-8') as f:
            content = f.read()
        
        # 简单的 bib 解析
        entries = re.findall(r'@(\w+)\{([^,]+),([\s\S]+?)\n\}', content)
        
        for entry_type, key, fields_text in entries:
            fields = {}
            
            # 提取字段
            field_matches = re.findall(r'(\w+)\s*=\s*\{([^}]+)\}', fields_text)
            for field_name, field_value in field_matches:
                fields[field_name.lower()] = field_value
            
            references.append({
                "key": key.strip(),
                "type": entry_type,
                "title": fields.get("title", ""),
                "authors": fields.get("author", ""),
                "year": fields.get("year", ""),
                "doi": fields.get("doi", "")
            })
    
    except Exception:
        pass
    
    return references
